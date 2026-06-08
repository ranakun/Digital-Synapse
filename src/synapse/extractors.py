"""Format-specific source extractors for ingestion.

The module keeps extraction deterministic and dependency-light. Optional
parsers are imported lazily so the package remains usable even when ingest
extras are not installed.
"""

from __future__ import annotations

import csv
import io
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import yaml

try:  # python-frontmatter is a core dependency, but keep a fallback.
    import frontmatter
except Exception:  # pragma: no cover - fallback is exercised only when missing.
    frontmatter = None  # type: ignore[assignment]

__all__ = [
    "ExtractedText",
    "extract",
    "extract_file",
    "extract_text",
    "extract_source_text",
    "extract_txt",
    "extract_md",
    "extract_csv",
    "extract_json",
    "extract_pdf",
    "extract_docx",
    "extract_vcard",
    "extract_whatsapp_chat",
]


_WHATSAPP_LINE_RE = re.compile(
    r"^(?:\u200e|\u200f)?(?P<ts>"
    r"(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4},?\s+\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM|am|pm)?)"
    r"|(?:\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}(?::\d{2})?)"
    r")\s*[-\u2013]\s*(?P<body>.*)$"
)
_WHATSAPP_SENDER_RE = re.compile(r"^(?P<sender>[^:]{1,120}):\s*(?P<message>.*)$")
_WHATSAPP_CONTINUATION_RE = re.compile(r"^\s{2,}\S")


@dataclass(frozen=True, slots=True)
class ExtractedText:
    """Normalized extracted text with format metadata."""

    text: str
    format_hint: str
    source_path: str | None = None
    source_format: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)
    warnings: tuple[str, ...] = ()

    @property
    def is_empty(self) -> bool:
        return not self.text.strip()


def extract_text(source: str | Path) -> ExtractedText:
    """Extract normalized text from a source file path."""

    path = Path(source)
    suffix = path.suffix.lower()
    name = path.name.lower()

    if suffix in {".md", ".markdown", ".mdown", ".mkdn"}:
        return extract_md(path)
    if suffix == ".csv":
        return extract_csv(path)
    if suffix == ".json":
        return extract_json(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".vcf", ".vcard"}:
        return extract_vcard(path)

    raw = _read_text_file(path)
    if _looks_like_whatsapp_export(raw, path.name):
        return extract_whatsapp_chat(path, raw_text=raw)
    if suffix in {".txt", ""} or name.endswith(".txt"):
        return extract_txt(path, raw_text=raw)

    return ExtractedText(
        text=_normalize_text(raw),
        format_hint="text",
        source_path=str(path),
        source_format=suffix.lstrip(".") or "unknown",
        metadata={},
        warnings=(f"unsupported extension: {suffix or '<none>'}",),
    )


def extract(source: str | Path) -> ExtractedText:
    return extract_text(source)


def extract_file(source: str | Path) -> ExtractedText:
    return extract_text(source)


def extract_source_text(source: str | Path) -> ExtractedText:
    return extract_text(source)


def extract_txt(path: str | Path, *, raw_text: str | None = None) -> ExtractedText:
    text = _read_text_file(Path(path)) if raw_text is None else raw_text
    return ExtractedText(
        text=_normalize_text(text),
        format_hint="text",
        source_path=str(Path(path)),
        source_format="txt",
        metadata={},
    )


def extract_md(path: str | Path) -> ExtractedText:
    text = _read_text_file(Path(path))
    body, metadata = _split_frontmatter(text)
    return ExtractedText(
        text=_normalize_text(body),
        format_hint="markdown",
        source_path=str(Path(path)),
        source_format="md",
        metadata=metadata,
    )


def extract_csv(path: str | Path) -> ExtractedText:
    raw = _read_text_file(Path(path))
    text, metadata = _parse_csv(raw)
    return ExtractedText(
        text=text,
        format_hint="csv",
        source_path=str(Path(path)),
        source_format="csv",
        metadata=metadata,
    )


def extract_json(path: str | Path) -> ExtractedText:
    raw = _read_text_file(Path(path))
    text, metadata = _parse_json(raw)
    return ExtractedText(
        text=text,
        format_hint="json",
        source_path=str(Path(path)),
        source_format="json",
        metadata=metadata,
    )


def extract_pdf(path: str | Path) -> ExtractedText:
    path = Path(path)
    try:
        import fitz  # type: ignore[import-not-found]
    except Exception as exc:  # pragma: no cover - depends on optional install.
        return ExtractedText(
            text="",
            format_hint="pdf",
            source_path=str(path),
            source_format="pdf",
            metadata={},
            warnings=(f"pdf extraction unavailable: {exc.__class__.__name__}",),
        )

    try:
        doc = fitz.open(str(path))
    except Exception as exc:
        return ExtractedText(
            text="",
            format_hint="pdf",
            source_path=str(path),
            source_format="pdf",
            metadata={},
            warnings=(f"unable to open pdf: {exc.__class__.__name__}",),
        )

    try:
        pages: list[str] = []
        for page in doc:
            pages.append(_normalize_text(page.get_text("text")))
        return ExtractedText(
            text="\n\n".join(page for page in pages if page.strip()),
            format_hint="pdf",
            source_path=str(path),
            source_format="pdf",
            metadata={"page_count": len(pages)},
        )
    finally:
        doc.close()


def extract_docx(path: str | Path) -> ExtractedText:
    path = Path(path)
    try:
        import docx  # type: ignore[import-not-found]
    except Exception as exc:  # pragma: no cover - depends on optional install.
        return ExtractedText(
            text="",
            format_hint="docx",
            source_path=str(path),
            source_format="docx",
            metadata={},
            warnings=(f"docx extraction unavailable: {exc.__class__.__name__}",),
        )

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        return ExtractedText(
            text="",
            format_hint="docx",
            source_path=str(path),
            source_format="docx",
            metadata={},
            warnings=(f"unable to open docx: {exc.__class__.__name__}",),
        )

    paragraphs = [_normalize_text(paragraph.text) for paragraph in document.paragraphs]
    tables = [_extract_docx_table(table) for table in document.tables]
    sections = [section for section in paragraphs + tables if section.strip()]
    return ExtractedText(
        text="\n\n".join(sections),
        format_hint="docx",
        source_path=str(path),
        source_format="docx",
        metadata={
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
        },
    )


def extract_vcard(path: str | Path) -> ExtractedText:
    path = Path(path)
    raw = _read_text_file(path)
    try:
        import vobject  # type: ignore[import-not-found]
    except Exception as exc:  # pragma: no cover - depends on optional install.
        text, metadata = _parse_vcard_fallback(raw)
        return ExtractedText(
            text=text,
            format_hint="vcard",
            source_path=str(path),
            source_format="vcard",
            metadata=metadata,
            warnings=(
                f"vcard extraction unavailable: {exc.__class__.__name__}; used fallback parser",
            ),
        )

    contacts: list[str] = []
    card_count = 0
    try:
        for card in vobject.readComponents(io.StringIO(raw)):
            card_count += 1
            lines: list[str] = []
            for label, field in (
                ("FN", "fn"),
                ("EMAIL", "email"),
                ("TEL", "tel"),
                ("ORG", "org"),
                ("TITLE", "title"),
                ("NOTE", "note"),
            ):
                value = _first_vcard_value(card, field)
                if value:
                    lines.append(f"{label}: {value}")
            contacts.append("\n".join(lines) if lines else _normalize_text(card.serialize()))
    except Exception as exc:
        return ExtractedText(
            text=_normalize_text(raw),
            format_hint="vcard",
            source_path=str(path),
            source_format="vcard",
            metadata={},
            warnings=(f"unable to parse vcard: {exc.__class__.__name__}",),
        )

    return ExtractedText(
        text="\n\n".join(section for section in contacts if section.strip()),
        format_hint="vcard",
        source_path=str(path),
        source_format="vcard",
        metadata={"card_count": card_count},
    )


def extract_whatsapp_chat(
    path: str | Path,
    *,
    raw_text: str | None = None,
) -> ExtractedText:
    text = _read_text_file(Path(path)) if raw_text is None else raw_text
    normalized, metadata = _parse_whatsapp_export(text)
    return ExtractedText(
        text=normalized,
        format_hint="whatsapp-chat",
        source_path=str(Path(path)),
        source_format="txt",
        metadata=metadata,
    )


def _read_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8-sig", errors="replace")
    except FileNotFoundError:
        raise
    except Exception:
        return path.read_bytes().decode("utf-8", errors="replace")


def _normalize_text(text: str) -> str:
    cleaned = text.replace("\r\n", "\n").replace("\r", "\n").replace("\ufeff", "")
    lines = [line.rstrip() for line in cleaned.split("\n")]
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    return "\n".join(lines)


def _split_frontmatter(text: str) -> tuple[str, dict[str, Any]]:
    if frontmatter is not None:
        try:
            post = frontmatter.loads(text)
            return str(post.content), dict(post.metadata or {})
        except Exception:
            pass

    if not text.startswith("---\n"):
        return text, {}

    parts = text.split("\n---\n", 1)
    if len(parts) != 2:
        return text, {}

    header = parts[0][4:]
    body = parts[1]
    try:
        metadata = yaml.safe_load(header) or {}
        if not isinstance(metadata, dict):
            metadata = {}
    except Exception:
        metadata = {}
    return body, metadata


def _parse_csv(raw: str) -> tuple[str, dict[str, Any]]:
    reader = csv.reader(io.StringIO(raw))
    rows = list(reader)
    if not rows:
        return "", {"row_count": 0, "columns": []}

    header = rows[0]
    data_rows = rows[1:]
    lines = ["COLUMNS: " + " | ".join(header)]
    for index, row in enumerate(data_rows, start=1):
        if not any(cell.strip() for cell in row):
            continue
        padded = list(row) + [""] * max(0, len(header) - len(row))
        values = [f"{column}={cell}" for column, cell in zip(header, padded, strict=False)]
        lines.append(f"ROW {index}: " + " | ".join(values))
    return "\n".join(lines), {"row_count": len(data_rows), "columns": header}


def _parse_json(raw: str) -> tuple[str, dict[str, Any]]:
    data = json.loads(raw)
    metadata = {"json_type": type(data).__name__}
    if isinstance(data, list):
        metadata["item_count"] = len(data)
    elif isinstance(data, dict):
        metadata["keys"] = list(data.keys())
    if _looks_like_linkedin_json(data):
        metadata["source_shape"] = "linkedin"
    return json.dumps(data, indent=2, sort_keys=True, ensure_ascii=False), metadata


def _looks_like_linkedin_json(data: Any) -> bool:
    if isinstance(data, dict):
        keys = {str(key).lower() for key in data}
        return bool(
            {"elements", "profile", "profiles", "positions", "firstname", "lastname"} & keys
        )
    if isinstance(data, list) and data:
        first = data[0]
        if isinstance(first, dict):
            keys = {str(key).lower() for key in first}
            return bool({"firstname", "lastname", "companyname", "position", "emailaddress"} & keys)
    return False


def _parse_whatsapp_export(raw: str) -> tuple[str, dict[str, Any]]:
    lines = _normalize_text(raw).splitlines()
    messages: list[str] = []
    participants: set[str] = set()
    current_message: list[str] = []

    for line in lines:
        if not line.strip():
            continue
        if _WHATSAPP_LINE_RE.match(line):
            if current_message:
                messages.append("\n".join(current_message).rstrip())
                current_message = []
            current_message.append(_normalize_whatsapp_line(line, participants))
        elif _WHATSAPP_CONTINUATION_RE.match(line) and current_message:
            current_message.append(line.strip())
        else:
            current_message.append(line.strip())

    if current_message:
        messages.append("\n".join(current_message).rstrip())

    return "\n\n".join(message for message in messages if message.strip()), {
        "message_count": len(messages),
        "participants": sorted(participants),
    }


def _normalize_whatsapp_line(line: str, participants: set[str]) -> str:
    match = _WHATSAPP_LINE_RE.match(line)
    if not match:
        return line.strip()

    body = match.group("body")
    sender_match = _WHATSAPP_SENDER_RE.match(body)
    if sender_match:
        sender = sender_match.group("sender").strip()
        participants.add(sender)
        return f"{match.group('ts')} - {sender}: {sender_match.group('message').strip()}"

    return f"{match.group('ts')} - {body.strip()}"


def _looks_like_whatsapp_export(text: str, filename: str) -> bool:
    lower_name = filename.lower()
    if lower_name.startswith("whatsapp") or "whatsapp" in lower_name:
        return True
    for line in text.splitlines()[:25]:
        if _WHATSAPP_LINE_RE.match(line):
            return True
    return False


def _extract_docx_table(table: Any) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [_normalize_text(cell.text) for cell in row.cells]
        if any(cell.strip() for cell in cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def _first_vcard_value(card: Any, field: str) -> str | None:
    component = None
    for attr in (field, field.upper()):
        try:
            component = getattr(card, attr)
            break
        except Exception:
            continue
    if component is None:
        return None
    values = getattr(component, "value", None)
    if values is not None:
        if isinstance(values, (list, tuple)):
            text = " ".join(str(item).strip() for item in values if str(item).strip())
        else:
            text = str(values).strip()
        return text or None
    try:
        text = str(component[0].value).strip()
        return text or None
    except Exception:
        return None


def _parse_vcard_fallback(raw: str) -> tuple[str, dict[str, Any]]:
    lines = _unfold_lines(_normalize_text(raw).splitlines())
    cards: list[dict[str, list[str]]] = []
    current: dict[str, list[str]] | None = None

    for line in lines:
        upper = line.upper()
        if upper == "BEGIN:VCARD":
            current = {}
            continue
        if upper == "END:VCARD":
            if current is not None:
                cards.append(current)
            current = None
            continue
        if current is None:
            continue
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        key = key.split(";", 1)[0].strip().upper()
        value = value.strip()
        current.setdefault(key, []).append(value)

    if current is not None:
        cards.append(current)

    rendered: list[str] = []
    for card in cards:
        lines_out: list[str] = []
        fn = _first_card_value(card, "FN") or _compose_vcard_name(card)
        if fn:
            lines_out.append(f"FN: {fn}")
        for label in ("EMAIL", "TEL", "ORG", "TITLE", "NOTE"):
            value = _first_card_value(card, label)
            if value:
                lines_out.append(f"{label}: {value}")
        if not lines_out:
            lines_out.extend(
                f"{key}: {values[0]}" for key, values in sorted(card.items()) if values
            )
        rendered.append("\n".join(lines_out))

    return "\n\n".join(section for section in rendered if section.strip()), {
        "card_count": len(cards),
        "parser": "fallback",
    }


def _first_card_value(card: dict[str, list[str]], key: str) -> str | None:
    values = card.get(key)
    if not values:
        return None
    text = values[0].strip()
    return text or None


def _compose_vcard_name(card: dict[str, list[str]]) -> str | None:
    values = card.get("N")
    if not values:
        return None
    parts = [part.strip() for part in values[0].split(";")]
    if not parts:
        return None
    given = parts[1] if len(parts) > 1 else ""
    family = parts[0] if parts else ""
    name = " ".join(part for part in [given, family] if part).strip()
    return name or None


def _unfold_lines(lines: list[str]) -> list[str]:
    unfolded: list[str] = []
    for line in lines:
        if line.startswith((" ", "\t")) and unfolded:
            unfolded[-1] += line[1:]
        else:
            unfolded.append(line)
    return unfolded
